import requests
import streamlit as st
from datetime import datetime, timezone
from transformers import pipeline
from docx import Document
from streamlit_option_menu import option_menu

#GitHub Confiq
GITHUB_TOKEN = "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"
GITHUB_REPO = "OwenAw/AIROST-Knowledge-Base"

# Notion API Config
NOTION_TOKEN = "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"
DATABASE_ID = "1494d66e48f480828afeed2595d52910"
headers = {
    "Authorization": f"Bearer {NOTION_TOKEN}",
    "Content-Type": "application/json",
    "Notion-Version": "2022-06-28",
}

# Load Hugging Face models
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
text_generator = pipeline("text-generation", model="gpt2")

#Functions
#Function to read files from github
def fetch_files_from_github():
    url = f"https://api.github.com/repos/{GITHUB_REPO}/contents"
    headers = {
        "Authorization": f"token {GITHUB_TOKEN}",
        "Accept": "application/vnd.github.v3+json"
    }
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        files = response.json()
        return files
    else:
        print(f"Error: {response.status_code}, {response.json()}")
        return []

#Function to displat GitHub Files
def display_files_in_dashboard(files):
    st.title("Project Files")
    for file in files:
        st.markdown(f"- [{file['name']}]({file['html_url']})")

#Function to create a new page in the Notion database    
def create_page(name, description, tags, github_link):
    create_url = "https://api.notion.com/v1/pages"
    payload = {
        "parent": {"database_id": DATABASE_ID},
        "properties": {
            "Name": {"title": [{"text": {"content": name}}]},
            "Description": {"rich_text": [{"text": {"content": description}}]},
            "Tags": {"multi_select": [{"name": tag} for tag in tags]},
            "GitHub Link": {"url": github_link},
        },
    }
    response = requests.post(create_url, headers=headers, json=payload)
    return response

#Function to fetch all projects from the Notion database
def fetch_pages():
    query_url = f"https://api.notion.com/v1/databases/{DATABASE_ID}/query"
    payload = {"page_size": 100}
    response = requests.post(query_url, headers=headers, json=payload)
    data = response.json()
    return data.get("results", [])

#Function to fetch filtered pages
def fetch_filtered_pages(search_text=None, tag_filter=None):
    url = f"https://api.notion.com/v1/databases/{DATABASE_ID}/query"
    filters = []

    if tag_filter:
        filters.append({
            "property": "Tags",
            "multi_select": {"contains": tag_filter}
        })

    if search_text:
        filters.append({
            "property": "Name",
            "title": {"contains": search_text}
        })

    payload = {"filter": {"and": filters}} if filters else {}
    response = requests.post(url, json=payload, headers=headers)
    
    if response.status_code != 200:
        print("Error fetching projects:", response.status_code, response.text)
        response.raise_for_status()
    
    results = response.json()["results"]
    projects = []
    for result in results:
        try:
            project_title = result["properties"]["Name"]["title"][0]["text"]["content"]
            project_tags = [tag["name"] for tag in result["properties"].get("Tags", {}).get("multi_select", [])]
            project_url = result["properties"].get("GitHub Link", {}).get("url", "No link")  # Fetch the project URL from the metadata
            projects.append({
                "title": project_title,
                "tags": project_tags,
                "url": project_url
            })
        except KeyError as e:
            print(f"Missing property: {e}")
    return projects

#Function to generate a summary
def generate_summary(description):
    summary = summarizer(description, max_length=50, min_length=10, do_sample=False)
    return summary[0]["summary_text"]

#Function to generate detailed documentation
def generate_documentation(prompt):
    response = text_generator(prompt, max_length=200, num_return_sequences=1)
    return response[0]["generated_text"]

#Function to create a Word document
def create_report(title, description, summary, detailed_doc):
    doc = Document()
    doc.add_heading(title, level=1)

    doc.add_heading("Project Description:", level=2)
    doc.add_paragraph(description)

    doc.add_heading("AI-Generated Summary:", level=2)
    doc.add_paragraph(summary)

    doc.add_heading("Detailed Documentation:", level=2)
    doc.add_paragraph(detailed_doc)

    # Save the document
    filename = f"{title.replace(' ', '_')}_Report.docx"
    doc.save(filename)
    return filename

#Streamlit App
# Sidebar Menu with Dropdown
with st.sidebar:
    selected = option_menu(
        menu_title="Navigation",  # Sidebar title
        options=["Welcome", "Code Repo", "Documentation", "Search and Filter", "AI Documentation Generator"],
        icons=["house", "github", "book", "search", "robot"],  # Optional icons from font-awesome
        menu_icon="cast",  # Main menu icon
        default_index=0,  # Default selected menu
        orientation="vertical",
    )

# Welcome Page
if selected == "Welcome":
    st.title("Welcome to the AIROST Knowledge Base")
    st.subheader("About This Platform")
    st.write("""
    This website is designed to streamline project management and documentation for AIROST teams. 
    Here’s what you can do:
    - Access the latest code repositories.
    - View detailed project documentation.
    - Use powerful search and filter tools to find specific projects.
    - Generate AI-assisted project documentation effortlessly.
    Navigate through the menu to explore different sections.
    """)

# Code Repo Section
if selected == "Code Repo":
    st.title("Code Repository")
    st.write("Here, you can view all the files in the GitHub repository.")
    files = fetch_files_from_github()
    if files:
        display_files_in_dashboard(files)
    else:
        st.warning("No files found or failed to fetch data.")

# Documentation Section
if selected == "Documentation":
    st.title("Documentation Viewer")
    st.write("View and uploade project documentation.")
    menu = st.sidebar.selectbox("Menu", ["View Projects", "Add Project"])
    if menu == "View Projects":
        st.subheader("All Projects")
        projects = fetch_pages()
        for project in projects:
            props = project["properties"]
            name = props["Name"]["title"][0]["text"]["content"]
            description = props.get("Description", {}).get("rich_text", [{}])[0].get("text", {}).get("content", "No description")
            tags = [tag["name"] for tag in props.get("Tags", {}).get("multi_select", [])]
            github_link = props.get("GitHub Link", {}).get("url", "No link")
            st.write(f"### {name}")
            st.write(f"**Description:** {description}")
            st.write(f"**Tags:** {', '.join(tags)}")
            st.write(f"**GitHub Link:** [{github_link}]({github_link})")

    elif menu == "Add Project":
        st.subheader("Add New Project")
        name = st.text_input("Project Name")
        description = st.text_area("Description")
        tags = st.text_input("Tags (comma-separated)").split(",")
        github_link = st.text_input("GitHub Link")
        if st.button("Add Project"):
            response = create_page(name, description, tags, github_link)
            if response.status_code == 200:
                st.success("Project added successfully!")
            else:
                st.error(f"Failed to add project: {response.text}")

# Search and Filter Section
if selected == "Search and Filter":
    st.title("Search and Filter Projects")
    st.write("Easily find projects using title or tag filters.")
    search_text = st.text_input("Search by Title")
    tag_filter = st.selectbox("Search by Tag:", ["AI", "Robotics", "IoT", ""])

    if st.button("Search"):
        try:
            projects = fetch_filtered_pages(search_text=search_text, tag_filter=tag_filter)
            if projects:
                for project in projects:
                    st.subheader(project["title"])
                    st.write(f"Tags: {', '.join(project['tags'])}")
                    st.markdown(f"[View Project]({project['url']})")  # Create a clickable URL link
                    st.write("---")
            else:
                st.write("No projects found.")
        except Exception as e:
            st.error(f"Error fetching projects: {e}")

# AI Documentation Generator Section
if selected == "AI Documentation Generator":
    st.title("AI-Powered Documentation Generator")
    st.write("Generate project summaries and detailed documentation with AI.")
    project_title = st.text_input("Enter Project Title:")
    project_description = st.text_area("Enter Project Description:")

    if st.button("Generate Report"):
        if project_title and project_description:
            # Generate summary and detailed documentation
            with st.spinner("Generating AI-Powered Summary..."):
                summary = generate_summary(project_description)
            
            with st.spinner("Generating Detailed Documentation..."):
                prompt = f"Generate a detailed documentation for the following project:\n{project_description}"
                detailed_doc = generate_documentation(prompt)
            
            # Create report
            with st.spinner("Creating Word Document..."):
                report_path = create_report(project_title, project_description, summary, detailed_doc)
            
            st.success(f"Report successfully generated: {report_path}")
            st.download_button(
                label="Download Report",
                data=open(report_path, "rb").read(),
                file_name=report_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Please enter both a project title and description.")