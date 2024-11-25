from transformers import pipeline
from docx import Document
import streamlit as st

# Load Hugging Face models
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
text_generator = pipeline("text-generation", model="gpt2")

# Function to generate a summary
def generate_summary(description):
    summary = summarizer(description, max_length=50, min_length=10, do_sample=False)
    return summary[0]["summary_text"]

# Function to generate detailed documentation
def generate_documentation(prompt):
    response = text_generator(prompt, max_length=200, num_return_sequences=1)
    return response[0]["generated_text"]

# Function to create a Word document
def create_report(title, description, summary, detailed_doc):
    doc = Document()
    doc.add_heading(title, level=1)

    doc.add_heading("Project Description:", level=2)
    doc.add_paragraph(description)

    doc.add_heading("AI-Generated Summary:", level=2)
    doc.add_paragraph(summary)

    doc.add_heading("Detailed Documentation:", level=2)
    doc.add_paragraph(detailed_doc)

    # Save the document
    filename = f"{title.replace(' ', '_')}_Report.docx"
    doc.save(filename)
    return filename

# Streamlit App
st.title("AI-Powered Documentation Generator")
st.write("Generate project summaries and detailed documentation with one click.")

# User Inputs
project_title = st.text_input("Enter Project Title:")
project_description = st.text_area("Enter Project Description:")

if st.button("Generate Report"):
    if project_title and project_description:
        # Generate summary and detailed documentation
        with st.spinner("Generating AI-Powered Summary..."):
            summary = generate_summary(project_description)
        
        with st.spinner("Generating Detailed Documentation..."):
            prompt = f"Generate a detailed documentation for the following project:\n{project_description}"
            detailed_doc = generate_documentation(prompt)
        
        # Create report
        with st.spinner("Creating Word Document..."):
            report_path = create_report(project_title, project_description, summary, detailed_doc)
        
        st.success(f"Report successfully generated: {report_path}")
        st.download_button(
            label="Download Report",
            data=open(report_path, "rb").read(),
            file_name=report_path,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("Please enter both a project title and description.")



